#!/usr/bin/env python3
"""Generate docs/execution-plan-next-steps.docx — the owner-facing printable
walkthrough of the immediate launch path (Tier 0-2).

This is a SNAPSHOT generator. TASKS_FOR_ERIC.md stays the source of truth;
re-run this whenever tiers/status change:

    pip install python-docx
    python scripts/build-execution-plan-doc.py

Revision 2026-07-04b: LLC registered; Google Workspace live (eric@ + greg@);
Meta Business access in hand (CAPI step now actionable); GL insurance in
progress (LSA still gated on it).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0x9A, 0x5F, 0x0F)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)
CREAM = "FBF3E0"

doc = Document()

# ---- base styles ---------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

doc.styles["Heading 1"].font.name = "Calibri"
doc.styles["Heading 1"].font.size = Pt(18)
doc.styles["Heading 1"].font.bold = True
doc.styles["Heading 1"].font.color.rgb = NAVY
doc.styles["Heading 2"].font.name = "Calibri"
doc.styles["Heading 2"].font.size = Pt(14)
doc.styles["Heading 2"].font.bold = True
doc.styles["Heading 2"].font.color.rgb = NAVY
doc.styles["Heading 3"].font.name = "Calibri"
doc.styles["Heading 3"].font.size = Pt(12)
doc.styles["Heading 3"].font.bold = True
doc.styles["Heading 3"].font.color.rgb = GOLD

for s in doc.sections:
    s.top_margin = Inches(1); s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.25); s.right_margin = Inches(1.25)


# ---- helpers -------------------------------------------------------------
def _run(p, text, bold=False, italic=False, color=None, name=None, size=None):
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if name:
        r.font.name = name
    if size:
        r.font.size = Pt(size)
    return r


def h(level, text):
    doc.add_heading(text, level=level)


def title_block():
    p = doc.add_paragraph()
    _run(p, "Northvale Roofing", bold=True, color=NAVY, size=26)
    p2 = doc.add_paragraph()
    _run(p2, "Your Next Steps — Execution Plan", bold=True, color=GOLD, size=15)
    p3 = doc.add_paragraph()
    _run(p3, "Immediate priorities (Tier 0–2) · Revenue-ASAP window · "
             "Prepared 2026-07-04 · Revised after LLC + Google Workspace + Meta access",
         color=GRAY, size=10)


def para(text):
    return doc.add_paragraph(text)


def lead(lead_text, rest):
    p = doc.add_paragraph()
    _run(p, lead_text, bold=True)
    _run(p, rest)
    return p


def why(text):
    p = doc.add_paragraph()
    _run(p, "Why it matters.  ", bold=True, italic=True, color=GRAY)
    _run(p, text, italic=True, color=GRAY)
    return p


def done(text):
    p = doc.add_paragraph()
    _run(p, "✓ Done when:  ", bold=True, color=GREEN)
    _run(p, text)
    return p


def num(text):
    return doc.add_paragraph(text, style="List Number")


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def var(name, desc):
    p = doc.add_paragraph()
    _run(p, name, bold=True, name="Consolas", color=GOLD, size=10)
    _run(p, "  — " + desc, size=10)
    return p


def _shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def callout(label, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    _shade(c, CREAM)
    p = c.paragraphs[0]
    _run(p, label + "  ", bold=True, color=NAVY)
    _run(p, text)
    doc.add_paragraph()
    return t


# =========================================================================
title_block()

# ---- Read this first -----------------------------------------------------
h(2, "Read this first")
para("The website and the entire lead pipeline (web form + tracked-call intake "
     "→ CRM contact + deal + instant alert) are already built and tested. They "
     "sit dormant, waiting on the account setups below. Most of these steps live "
     "outside the code — in Google, HighLevel, Meta, Vercel, and your state "
     "filing. Once you paste each key into Vercel, the matching feature switches "
     "on at the next deploy.")
lead("The one big idea. ",
     "The website is a 6–12 month SEO asset — it will not ring the phone this "
     "quarter. What produces near-term revenue is the Tier 0 block: Google "
     "Business Profile + Local Services Ads + review velocity + sub-5-minute "
     "lead response. Do Tier 0 first, even before the SEO pages matter.")
lead("Pasting a key into Vercel (you'll do this a lot). ",
     "Go to Vercel → your project → Settings → Environment Variables → add the "
     "name and value (set it for Production and Preview) → Save. Vercel "
     "redeploys automatically. That's the whole loop every time you see a "
     "VAR_NAME below.")

callout("Ground rules:",
        "Claude never touches secrets, billing, DNS, or auth. Anything involving "
        "those is on this list because it is your action. Never fabricate "
        "reviews, addresses, or stats — real data only.")

# ---- Where you stand today (new) ----------------------------------------
h(2, "Where you stand today")
para("Progress since the first version of this plan:")
bullet("✓ LLC registered — Northvale Roofing LLC is filed. This unblocks the LSA "
       "application, insurance, the trademark filing, and the mailing address "
       "for legal pages.")
bullet("✓ Business email live — Google Workspace with eric@northvaleroofing.com "
       "and greg@northvaleroofing.com. (One small follow-up below: a shared "
       "hello@ alias + confirming the SPF/DKIM/DMARC records.)")
bullet("✓ Meta Business access — you now have the account, so the Conversions "
       "API step (Tier 2, Step 5) is actionable today instead of blocked.")
bullet("⏳ General-liability insurance — in progress (a few weeks out). This "
       "gates exactly ONE thing: the Local Services Ads application (Tier 0, "
       "Step 2). Everything else below can be done now.")

para("Do-today shortlist (none of these need insurance): start GBP verification "
     "(Tier 0, Step 1); stand up HighLevel + the two speed-to-lead workflows "
     "(Tier 0, Step 3 / Tier 2, Step 1); turn on lead alerts + Blob storage "
     "(Tier 2, Step 0); wire GA4 (Tier 2, Step 3); issue the Meta Pixel + CAPI "
     "token (Tier 2, Step 5); and set PUBLIC_SITE_URL in Vercel (Tier 1, "
     "Step 3). Then hand the facts list back to Claude to publish the gated pages.")

# ---- At a glance ---------------------------------------------------------
h(2, "At a glance")
tbl = doc.add_table(rows=4, cols=3)
tbl.style = "Table Grid"
rows = [
    ("Tier", "What it unlocks", "Status / blocks"),
    ("Tier 0", "Phone-ringers & speed-to-lead (GBP, LSA, HighLevel auto-response, reviews)",
     "Near-term revenue — do first (LSA waits on insurance)"),
    ("Tier 1", "Foundation to go live (LLC, email, phone number)",
     "LLC ✓ · email ✓ — one Vercel value left"),
    ("Tier 2", "Analytics + CRM wiring (HighLevel, CallRail, GA4, Meta, GBP Place ID)",
     "Blocks attribution / ROI — all actionable now"),
]
for ri, row in enumerate(rows):
    for ci, val in enumerate(row):
        cell = tbl.rows[ri].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        _run(p, val, bold=(ri == 0), color=(NAVY if ri == 0 else None))
        if ri == 0:
            _shade(cell, CREAM)
doc.add_paragraph()

# =========================================================================
# TIER 0
# =========================================================================
h(1, "Tier 0 — Phone-ringers & speed-to-lead")
para("Fastest path to a ringing phone. Because both owners are part-time, the "
     "automated pieces matter most — a lead can't wait on a busy person. These "
     "must be live before you spend the first ad dollar.")

h(3, "Step 1 — Claim & verify your Google Business Profile (GBP)   ▶ do today")
why("The local 3-pack and Local Services Ads sit above organic search. This is "
    "the single biggest near-term lever, and verification can take days — start "
    "the clock now.")
num("Go to business.google.com and sign in with a business Google account — "
    "eric@northvaleroofing.com works, or create a shared hello@ alias/group so "
    "the profile isn't tied to one person.")
num("Create the profile for 'Northvale Roofing'. Choose the roofing-contractor "
    "category. Northvale is a service-area business, so set your service areas "
    "(the 8 cities) rather than a storefront address.")
num("Complete verification (Google will offer phone, video, or postcard). "
    "Verification can take a few days — start it early.")
num("Fill the profile fully: services, hours, service areas, description, and "
    "add real photos as they exist.")
num("Copy your public GBP URL and paste it into the site admin panel "
    "(/keystatic → Business Info → profile URLs). This feeds the sameAs "
    "structured data on every page.")
callout("Note:", "You'll also need the numeric Place ID later (Tier 2, Step 7). "
        "Grab it once the profile is verified.")
done("The profile shows 'Verified' and its URL is pasted into the admin panel.")

h(3, "Step 2 — Apply for Google Local Services Ads (LSA)   ⏳ waiting on insurance")
why("Roofing LSA cost-per-lead runs about $55–90 (roughly half of blended "
    "Google Ads CPL) and earns the 'Google Screened' badge that sits at the "
    "very top of results.")
num("Apply the day your general-liability insurance binds — the LLC is already "
    "registered, so insurance is the only remaining prerequisite.")
num("Go to ads.google.com/local-services-ads and start the roofing application "
    "for the Houston metro.")
num("Complete the background check and upload proof of insurance and licensing "
    "when prompted.")
num("Set your weekly budget once approved. Leads are charged per valid lead, "
    "not per click.")
callout("Prerequisite:", "Registered LLC ✓ + active general-liability insurance "
        "(in progress). You can't finish the application without the insurance. "
        "Tip: check whether the background-check step can be started ahead of "
        "the insurance binding so it's not the long pole later.")
done("The application is submitted and the background/insurance checks are in "
     "progress.")

h(3, "Step 3 — Turn on speed-to-lead in HighLevel   ▶ do today (after Tier 2, Step 1)")
why("Response-time studies put the biggest drop-off in the first five minutes. "
    "For two part-time owners the fix is automation, not discipline. This is the "
    "#1 operational priority.")
para("Requires the HighLevel account + keys from Tier 2, Step 1, plus a "
     "HighLevel phone number for SMS. Then build two workflows inside HighLevel "
     "(this is configuration, not code):")
bullet("Workflow 1 — Instant auto-responder SMS (under 60 seconds, every form + call lead):")
bullet("Trigger: new contact created / inbound webhook (the site already "
       "delivers form and tracked-call leads into HighLevel).")
bullet("Action: immediate SMS from the business number — e.g. \"Hi "
       "{{contact.first_name}}, this is Northvale Roofing — thanks for reaching "
       "out. A real person will call you shortly. Roof leaking right now? Call "
       "us at (713) 449-7661. Reply STOP to opt out.\"")
bullet("Workflow 2 — Missed-call text-back:")
bullet("Trigger: inbound call to the tracked/business number that is not answered.")
bullet("Action: auto-SMS within seconds — \"Sorry we missed your call — this is "
       "Northvale Roofing. What's going on with your roof? Text us here and "
       "we'll get right back to you. Reply STOP to opt out.\"")
bullet("This is the highest-ROI single automation for part-time owners: it "
       "converts a missed call into a text thread instead of losing it to the "
       "next contractor.")
callout("TCPA:", "The website already captures consent + timestamp + IP on every "
        "form lead, so automated SMS to a consenting lead is compliant. Keep a "
        "'Reply STOP' opt-out in every message. Full spec: "
        "docs/runbooks/speed-to-lead.md.")
done("A test form lead triggers the auto-responder in under 60s, and a test "
     "ring-out triggers the missed-call text-back — both with a working opt-out.")

h(3, "Step 4 — Run the review-request SOP starting at job #1")
why("Review velocity is the wedge that moves you up the GBP 3-pack. Start at "
    "the very first completed job, not later.")
num("When a job is marked complete (final walk-through done), ask the happy "
    "customer for a review in person — the single biggest lever.")
num("Send the request by SMS + email with a direct Google review link (you'll "
    "have the short-link once GBP is verified). Automate this via HighLevel "
    "once connected.")
num("Respond to every review, positive and negative, professionally.")
callout("Hard line:", "Real reviews only. Never write, buy, or incentivize "
        "reviews, and ask ALL completed customers, not just the happy ones — "
        "fabricating or gating reviews violates the FTC 2024 rule and Google "
        "policy. Full SOP: docs/sops/review-request.md.")
done("The review ask is part of your job-close routine and the first genuine "
     "reviews are landing on the profile.")

h(3, "Step 5 — Add lead-qualification fields + follow-up cadences in HighLevel")
why("Cheap to set up now, expensive to retrofit. Different lead types need "
    "different sales motions, and 20–30% of signable jobs are lost to "
    "post-estimate silence.")
bullet("Qualification fields: capture retail vs. insurance-claim, roof age, and "
       "timeline (emergency / weeks / just researching) so leads branch into "
       "the right nurture track.")
bullet("No-answer cadence: attempt contact, then automated SMS/email nudges; "
       "after 3 no-answers, drop into a 90-day monthly nurture rather than "
       "marking the lead dead.")
bullet("Quote-follow-up sequence: after an estimate is sent, a timed sequence "
       "(day 1, 3, 7, 14) — this directly recovers the post-estimate-silence "
       "losses.")
done("Insurance-claim leads, retail leads, and 'just researching' leads each "
     "land in a distinct nurture track, and no lead can silently die.")

# =========================================================================
# TIER 1
# =========================================================================
h(1, "Tier 1 — Foundation to go live")
para("These block anything going public. Most are now done (kept here so the "
     "picture is complete).")

h(3, "Step 1 — Register 'Northvale Roofing LLC'   ✓ done")
para("The LLC is registered. One follow-up: once the registered-agent / mailing "
     "address is final, hand it to the next Claude session so it can fill "
     "emailFooterAddress in the legal constants (CAN-SPAM requires a physical "
     "address in marketing email) and complete the address in the schema.")
done("The LLC is registered ✓ — remaining: hand the mailing address to Claude.")

h(3, "Step 2 — Domain   ✓ done")
para("northvaleroofing.com is registered at Squarespace Domains (2026-04-23), "
    "DNS managed there. No action needed unless DNS records change.")

h(3, "Step 3 — Vercel project   ▶ one value to set today")
para("The project exists and preview builds auto-deploy on every push. "
     "Remaining action:")
var("PUBLIC_SITE_URL", "set to https://northvaleroofing.com in Vercel env once "
    "you're ready to point at the real domain.")
done("PUBLIC_SITE_URL is set to the apex domain in Vercel.")

h(3, "Step 4 — Business email   ✓ done (small follow-ups)")
para("Google Workspace is live with eric@northvaleroofing.com and "
     "greg@northvaleroofing.com. Two quick follow-ups:")
num("Add a shared hello@northvaleroofing.com alias or group for public-facing "
    "use (GBP, LSA, lead alerts) so the brand isn't tied to one person's inbox.")
num("Confirm the MX + SPF, DKIM, and DMARC records validate in Squarespace's "
    "DNS manager, so marketing email lands and stays CAN-SPAM compliant. (DNS "
    "is your action — Claude will not touch it.)")
done("A shared hello@ address exists and SPF/DKIM/DMARC validate.")

h(3, "Step 5 — Dedicated business phone number   (optional — a line already works)")
why("A single, consistent number is the backbone of NAP consistency, GBP, and "
    "call tracking.")
para("The real business phone (713) 449-7661 is already live in the site data "
     "and satisfies the NAP index gate. This step is only about a "
     "tracked/dedicated line if you choose a separate one.")
num("If you provision a dedicated number (Twilio or Google Voice Business), "
    "hand it to the next Claude session to update BRAND.phoneE164 + phoneDisplay "
    "in src/lib/brand.ts — never hardcode it elsewhere.")
done("The number is chosen and reflected in brand.ts everywhere the site shows "
     "a phone.")

# =========================================================================
# TIER 2
# =========================================================================
h(1, "Tier 2 — Analytics + CRM wiring")
para("These unlock attribution and ROI reporting. Every one is env-gated: paste "
     "the key into Vercel and the feature turns on — no key means a clean skip, "
     "nothing breaks. All of Tier 2 is actionable today.")

h(3, "Step 0 — Day-one lead alerts + never-lose-a-lead storage   ▶ do today (quick win)")
why("Your instant safety net BEFORE HighLevel is even connected: every form "
    "submission texts and emails you within seconds, and every lead is saved to "
    "permanent private storage. Do this early — it's ~15 minutes total.")
p = doc.add_paragraph(); _run(p, "Text alerts (Twilio, ~$1/mo + ~1¢/text):", bold=True)
var("TWILIO_ACCOUNT_SID", "Account SID (starts AC…)")
var("TWILIO_AUTH_TOKEN", "Auth Token")
var("TWILIO_FROM", "the Twilio number you buy, in +1… format")
var("LEAD_ALERT_SMS_TO", "owner cell(s), comma-separated E.164, e.g. +12815550111,+18325550222")
p = doc.add_paragraph(); _run(p, "Email alerts (Resend, free to start):", bold=True)
var("RESEND_API_KEY", "the key (starts re_…)")
var("LEAD_ALERT_EMAIL_TO", "owner email(s), comma-separated — e.g. eric@ + greg@northvaleroofing.com")
var("LEAD_ALERT_EMAIL_FROM", "leads@resend.dev to start, or a verified domain sender")
p = doc.add_paragraph(); _run(p, "Never-lose-a-lead storage (Vercel Blob, ~2 min):", bold=True)
bullet("In Vercel → Storage → Create Database → Blob, create a store with "
       "Private access (leads contain customer PII).")
bullet("Connect it to the project. Vercel adds BLOB_READ_WRITE_TOKEN "
       "automatically — there is no key to copy. Redeploy.")
callout("Note:", "Full walkthrough: docs/setup-leads.md. Each channel turns on "
        "independently when its keys are present.")
done("A test form submission texts and emails you within seconds, and a JSON "
     "record appears in the Blob store.")

h(3, "Step 1 — HighLevel CRM keys   ▶ do today (powers Tier 0 speed-to-lead)")
why("Once connected, every form submission and tracked call auto-creates a "
    "contact and opens a deal, tagged by source. This is also the engine behind "
    "the Tier 0 speed-to-lead workflows.")
num("Create (or open) the Northvale Roofing sub-account. Settings → Business "
    "Profile shows the Location/Business ID (also the long id in the dashboard URL).")
num("Settings → Private Integrations → Create new integration. Grant Contacts + "
    "Opportunities scopes (read + write). Copy the token (shown once).")
num("Optional but recommended: Opportunities → Pipelines → Create Pipeline "
    "(e.g. New Lead → Contacted → Estimate → Won/Lost). Get the pipeline + "
    "first-stage IDs from the URL or the pipelines API.")
num("Create the contact custom fields the site writes to (Settings → Custom "
    "Fields), keys exactly: lead_source, service, notes, utm_source, "
    "utm_medium, utm_campaign, utm_content, utm_term, gclid, fbclid, "
    "landing_path, first_touch_at (calls also use channel, callrail_call_id, "
    "tracking_number, call_duration, call_recording, callrail_source).")
num("Paste the values into Vercel:")
var("HIGHLEVEL_API_KEY", "the Private Integration token (required)")
var("HIGHLEVEL_LOCATION_ID", "the sub-account/location id (required)")
var("HIGHLEVEL_PIPELINE_ID", "optional — set with the stage below to also open a pipeline deal per lead")
var("HIGHLEVEL_PIPELINE_STAGE_ID", "optional — the first / 'New Lead' stage id")
callout("Note:", "First two keys = contacts created. Add the last two = each "
        "lead also opens an opportunity. Full guide: docs/setup-highlevel.md. "
        "After keys land, submit one test lead and confirm the live v2 payload "
        "shapes match — they were built to the published docs.")
done("A test form submission shows a new contact (and, if configured, a "
     "New-Lead opportunity) in HighLevel within seconds, with lead_source "
     "populated.")

h(3, "Step 2 — CallRail Elite (call tracking + attribution)")
why("Ties phone leads to the channel that produced them, and pushes completed "
    "calls into HighLevel with the same instant alert as form leads.")
num("Provision 3 tracking numbers: a site DNI pool, an offline pool, and a GBP pool.")
num("Point the Post-Call webhook at /api/callrail-webhook, secured with an "
    "HMAC-SHA1 secret you generate in CallRail → Webhooks.")
num("Paste the values into Vercel:")
var("CALLRAIL_WEBHOOK_SECRET", "the shared secret for webhook signature verification")
var("CALLRAIL_API_KEY", "CallRail API v3 key")
var("CALLRAIL_ACCOUNT_ID", "your CallRail account id")
var("PUBLIC_CALLRAIL_COMPANY_ID", "the companyId/scriptKey segments from the DNI swap.js snippet")
done("A completed test call to a tracked number appears as a HighLevel contact "
     "and fires the same SMS/email alert as a form lead.")

h(3, "Step 3 — Google Analytics 4 (GA4)   ▶ do today")
why("Baseline traffic + conversion measurement. The snippet is already wired "
    "and fires once the ID is present.")
num("Create a GA4 property at analytics.google.com → Admin → Create property → "
    "add a Web data stream for northvaleroofing.com.")
num("Copy the Measurement ID (G-XXXXXXXXXX) and set it in Vercel:")
var("PUBLIC_GA4_ID", "the G-XXXXXXXXXX measurement id")
callout("Bonus (2 min, no DNS):", "Verify Google Search Console via the HTML-tag "
        "method — set PUBLIC_GSC_VERIFICATION to the content value Google shows, "
        "redeploy, click Verify, then submit "
        "https://northvaleroofing.com/sitemap-index.xml. See "
        "docs/analytics-search-console.md.")
done("GA4 Realtime shows your own visit to the live site.")

h(3, "Step 4 — Server-side Google Tag Manager (optional)")
why("A server-side container improves measurement durability. Wired but only "
    "loads when the ID is set.")
num("Stand up a server-side GTM container on a Vercel subdomain (e.g. "
    "metrics.northvaleroofing.com).")
num("Set the ID in Vercel:")
var("PUBLIC_GTM_ID", "the GTM-XXXXXXX container id")
done("The GTM container is live and PUBLIC_GTM_ID is set.")

h(3, "Step 5 — Meta Business Manager + Conversions API (CAPI)   ▶ do today (you have access)")
why("Every form lead sends a server-side 'Lead' conversion (hashed PII) to Meta "
    "automatically — the code is done and waiting on keys. You now have the Meta "
    "account, so this is actionable today.")
num("In your Meta Business account, confirm/create the ad account, then open "
    "Events Manager → Data Sources and create a Pixel if one doesn't exist.")
num("Copy the numeric Pixel ID, then generate a Conversions API access token "
    "(Events Manager → your pixel → Settings → Generate access token).")
num("Set both in Vercel:")
var("META_CAPI_TOKEN", "the Conversions API access token (server-side)")
var("META_PIXEL_ID", "the pixel id (digits) — also renders the browser pixel")
callout("Note:", "Hold actual campaigns paused until the CPL signal is real "
        "(that's Tier 5). This step is just the account + pixel + token wiring.")
done("A test form lead shows a server-side Lead event in Meta Events Manager.")

h(3, "Step 6 — Google Ads account + conversion import")
why("Needed to run brand/non-brand search campaigns and import conversions from "
    "HighLevel for true ROI.")
num("Create a Google Ads account and set up a nightly conversion import from "
    "HighLevel.")
num("Set the conversion id (and label) in Vercel:")
var("GOOGLE_ADS_CONVERSION_ID", "AW-XXXXXXXXX")
var("GOOGLE_ADS_CONVERSION_LABEL", "the label after the slash — /thank-you "
    "fires the conversion only when both are set")
done("The account exists and the conversion id/label are set in Vercel.")

h(3, "Step 7 — Google Business Profile Place ID")
why("The numeric Place ID powers structured data and the live reviews pull. You "
    "get it once the GBP (Tier 0, Step 1) is verified.")
num("Find your Place ID with Google's Place ID finder "
    "(developers.google.com/maps/documentation/places/web-service/place-id).")
num("Set it in Vercel:")
var("GOOGLE_PLACE_ID", "the Northvale Roofing GBP Place ID")
done("GOOGLE_PLACE_ID is set in Vercel.")

h(3, "Step 8 — Google Places API key (live reviews pull)")
why("Lets /reviews pull real Google reviews. Real reviews unlock the "
    "AggregateRating schema.")
num("Create/scope a Google Places API key in Google Cloud for the reviews pull.")
num("Set it in Vercel:")
var("GOOGLE_PLACES_API_KEY", "scoped for the reviews pull")
done("GOOGLE_PLACES_API_KEY is set and /reviews pulls real GBP reviews.")

# ---- closing -------------------------------------------------------------
h(2, "How to hand work back to the site")
bullet("Keys → Vercel. Anything that's a KEY goes into Vercel env (Production + "
       "Preview). The feature activates on redeploy. Nothing else needed.")
bullet("Data → admin panel. GBP/social URLs, financing terms, team, warranty, "
       "and photos go into the /keystatic admin panel — those pages publish "
       "themselves once real data is present.")
bullet("Facts/decisions → next Claude session. The LLC's mailing address, years "
       "in market, carriers you've worked, crew model, and any new phone number "
       "— hand these to the next Claude session so it can update the code "
       "truthfully. Never invent any of them.")

h(2, "After Tier 0–2")
para("Once these are live and leads are flowing with attribution, the plan "
     "continues into Tier 2.5/2.6 (data that unlocks gated SEO pages + the "
     "copy-facts questionnaire), Tier 3 (legal review + real photography, "
     "removing demo images), and Tier 5 (paid campaigns + the launch "
     "checklist). The full ordered list lives in TASKS_FOR_ERIC.md in the repo. "
     "Ask for the next installment of this document when you're ready.")

p = doc.add_paragraph()
_run(p, "Northvale Roofing · Internal execution plan · Not for public "
        "distribution", color=GRAY, size=10)

doc.save("docs/execution-plan-next-steps.docx")
print("wrote docs/execution-plan-next-steps.docx")
